"""Markdown to DOCX converter using AST-based approach with mistune and python-docx."""

from docx import Document
from docx.document import Document as DocumentClass
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import mistune
from typing import Any


class DocxEmitter:
    """Emits Word document elements from a markdown AST."""

    def __init__(self) -> None:
        self.doc = Document()
        self._current_paragraph: Any = None

    def _ensure_paragraph(self) -> Any:
        """Ensure we have a current paragraph to add runs to."""
        if self._current_paragraph is None:
            self._current_paragraph = self.doc.add_paragraph()
        return self._current_paragraph

    def _finish_paragraph(self) -> None:
        """Mark current paragraph as finished."""
        self._current_paragraph = None

    def emit(self, tokens: list[dict]) -> None:
        """Process a list of AST tokens and emit document elements."""
        for token in tokens:
            self._emit_token(token)

    def _emit_token(self, token: dict) -> None:
        """Dispatch a single token to the appropriate handler."""
        token_type = token.get("type", "")
        handler = getattr(self, f"_emit_{token_type}", None)
        if handler:
            handler(token)
        elif "children" in token:
            # Fallback: process children for unknown container types
            self.emit(token["children"])

    def _emit_blank_line(self, token: dict) -> None:
        """Handle blank lines."""
        pass  # Ignore blank lines

    def _emit_paragraph(self, token: dict) -> None:
        """Handle paragraph elements."""
        self._finish_paragraph()
        p = self._ensure_paragraph()
        self._emit_inline(token.get("children", []), p)
        self._finish_paragraph()

    def _emit_heading(self, token: dict) -> None:
        """Handle headings (h1-h6)."""
        self._finish_paragraph()
        level = token.get("attrs", {}).get("level", 1)
        text = self._extract_text(token.get("children", []))
        self.doc.add_heading(text, level=level)
        self._finish_paragraph()

    def _emit_block_code(self, token: dict) -> None:
        """Handle code blocks."""
        self._finish_paragraph()
        p = self.doc.add_paragraph(style="No Spacing")
        code = token.get("raw", "")
        run = p.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        self._finish_paragraph()

    def _emit_block_quote(self, token: dict) -> None:
        """Handle blockquotes."""
        self._finish_paragraph()
        text = self._extract_text(token.get("children", []))
        p = self.doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.5)
        self._finish_paragraph()

    def _emit_thematic_break(self, token: dict) -> None:
        """Handle horizontal rules."""
        self._finish_paragraph()
        p = self.doc.add_paragraph("─" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._finish_paragraph()

    def _emit_list(self, token: dict) -> None:
        """Handle lists."""
        self._finish_paragraph()
        ordered = token.get("attrs", {}).get("ordered", False)
        style = "List Number" if ordered else "List Bullet"
        for child in token.get("children", []):
            if child.get("type") == "list_item":
                text = self._extract_text(child.get("children", []))
                self.doc.add_paragraph(text, style=style)
        self._finish_paragraph()

    def _emit_table(self, token: dict) -> None:
        """Handle tables by extracting data and creating Word table."""
        self._finish_paragraph()

        rows: list[list[str]] = []
        children = token.get("children", [])

        for child in children:
            if child.get("type") == "table_head":
                # Header cells are direct children of table_head
                header_row = []
                for cell in child.get("children", []):
                    if cell.get("type") == "table_cell":
                        header_row.append(self._extract_text(cell.get("children", [])))
                if header_row:
                    rows.append(header_row)
            elif child.get("type") == "table_body":
                # Body has table_row children, each with table_cell children
                for row in child.get("children", []):
                    if row.get("type") == "table_row":
                        row_data = []
                        for cell in row.get("children", []):
                            if cell.get("type") == "table_cell":
                                row_data.append(
                                    self._extract_text(cell.get("children", []))
                                )
                        if row_data:
                            rows.append(row_data)

        if not rows:
            return

        num_rows = len(rows)
        num_cols = max(len(row) for row in rows) if rows else 0

        if num_cols == 0:
            return

        # Create the table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        # Populate the table
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    cell.text = cell_text.strip()
                    # Bold the header row
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        # Add shading to header row
        if num_rows > 0:
            for cell in table.rows[0].cells:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)

        self._finish_paragraph()

    def _emit_inline(self, tokens: list[dict], paragraph: Any) -> None:
        """Emit inline elements to a paragraph."""
        for token in tokens:
            token_type = token.get("type", "")

            if token_type == "text":
                paragraph.add_run(token.get("raw", ""))
            elif token_type == "strong":
                text = self._extract_text(token.get("children", []))
                run = paragraph.add_run(text)
                run.bold = True
            elif token_type == "emphasis":
                text = self._extract_text(token.get("children", []))
                run = paragraph.add_run(text)
                run.italic = True
            elif token_type == "codespan":
                run = paragraph.add_run(token.get("raw", ""))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            elif token_type == "link":
                text = self._extract_text(token.get("children", []))
                url = token.get("attrs", {}).get("url", "")
                paragraph.add_run(f"{text} ({url})")
            elif token_type == "strikethrough":
                text = self._extract_text(token.get("children", []))
                run = paragraph.add_run(text)
                run.font.strike = True
            elif token_type == "linebreak":
                paragraph.add_run("\n")
            elif token_type == "softbreak":
                paragraph.add_run(" ")
            elif "children" in token:
                self._emit_inline(token["children"], paragraph)

    def _extract_text(self, tokens: list[dict]) -> str:
        """Recursively extract plain text from a list of tokens."""
        parts: list[str] = []
        for token in tokens:
            token_type = token.get("type", "")
            if token_type == "text":
                parts.append(token.get("raw", ""))
            elif token_type == "codespan":
                parts.append(token.get("raw", ""))
            elif token_type == "softbreak":
                parts.append(" ")
            elif token_type == "linebreak":
                parts.append("\n")
            elif "children" in token:
                parts.append(self._extract_text(token["children"]))
            elif "raw" in token:
                parts.append(token["raw"])
        return "".join(parts)


def convert_markdown_to_docx(markdown_text: str) -> DocumentClass:
    """Convert markdown text to a Word document.

    Args:
        markdown_text: The GitHub-flavored markdown text to convert.

    Returns:
        A python-docx Document object.
    """
    # Parse markdown to AST (renderer=None returns AST)
    md = mistune.create_markdown(
        renderer=None,
        plugins=["table", "strikethrough"],
    )
    tokens = md(markdown_text)

    # Emit document from AST
    emitter = DocxEmitter()
    if isinstance(tokens, list):
        emitter.emit(tokens)
    return emitter.doc


def convert_file(input_path: str, output_path: str) -> None:
    """Convert a markdown file to a Word document.

    Args:
        input_path: Path to the input markdown file.
        output_path: Path for the output .docx file.
    """
    with open(input_path, "r", encoding="utf-8") as f:
        markdown_text = f.read()

    doc = convert_markdown_to_docx(markdown_text)
    doc.save(output_path)
